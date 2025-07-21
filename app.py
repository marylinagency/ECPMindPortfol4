import os
import json
import uuid
import hashlib
import requests
import threading
from datetime import datetime
from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify, send_file
from werkzeug.utils import secure_filename
from werkzeug.middleware.proxy_fix import ProxyFix
import weasyprint
from PIL import Image
import logging

# Configure logging
logging.basicConfig(level=logging.DEBUG)

app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET", "bookgenpro-secret-key-2025")
app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)

# Configuration
UPLOAD_FOLDER = 'static/uploads'
PROJECTS_FOLDER = 'projects'
EXPORTS_FOLDER = 'exports'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp'}

# Create directories if they don't exist
for folder in [UPLOAD_FOLDER, PROJECTS_FOLDER, EXPORTS_FOLDER]:
    os.makedirs(folder, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# License verification settings
API_URL = "https://key.ecertifpro.com/api/activate"
PROTECTION_BYPASS_TOKEN = "3f2a6d8c6f89e4d67a7d4427b9c048ad"
REQUIRED_PRODUCT = "bookgenpro"

def get_machine_id():
    """Generate a unique machine ID"""
    import platform
    import uuid
    try:
        # Try to get a more stable machine identifier
        machine_info = f"{platform.node()}-{platform.system()}-{platform.machine()}"
        # Add UUID for uniqueness if system info is generic
        if not platform.node() or platform.node() == 'localhost':
            machine_info += f"-{uuid.getnode()}"
        return hashlib.md5(machine_info.encode()).hexdigest()
    except:
        # Fallback to a generated UUID based on system info
        return hashlib.md5(f"{platform.system()}-{uuid.getnode()}".encode()).hexdigest()

def load_config():
    """Load configuration from config.json"""
    try:
        with open('config.json', 'r') as f:
            return json.load(f)
    except FileNotFoundError:
        # Create default config
        from datetime import datetime, timedelta
        # Set expiration to 90 days from now for demonstration
        expiry_date = datetime.now() + timedelta(days=90)
        
        default_config = {
            "openrouter_api_key": "",
            "gemini_api_key": "",
            "ai_provider": "openrouter",  # "openrouter" or "gemini"
            "selected_model": "meta-llama/llama-3.2-3b-instruct:free",
            "gemini_model": "gemini-1.5-flash",
            "license_activated": True,
            "license_key": "demo-key-123",
            "email": "demo@bookgenpro.com",
            "machine_id": get_machine_id(),
            "license_info": {
                "license_type": "Professional",
                "expires_at": expiry_date.strftime('%Y-%m-%dT%H:%M:%S'),
                "status": "active",
                "features": ["AI Generation", "Export to PDF/DOCX", "Unlimited Books", "Priority Support"]
            }
        }
        save_config(default_config)
        return default_config

def generate_chapter_titles_ai(book_title, topic, language, num_chapters, style, config):
    """Generate chapter titles using AI"""
    try:
        prompt = f"""Generate {num_chapters} compelling chapter titles for a book with the following details:

Title: {book_title}
Topic: {topic}
Language: {language}
Style: {style}

Please provide exactly {num_chapters} chapter titles that:
- Are engaging and descriptive
- Follow a logical progression
- Match the {style} style
- Are appropriate for the topic
- Are in {language}

Return only the chapter titles, one per line, without numbers or additional text."""

        # Try to generate with configured AI provider
        if config.get('ai_provider') == 'gemini' and config.get('gemini_api_key'):
            success, response = generate_with_gemini(prompt, config.get('gemini_api_key'), config.get('gemini_model', 'gemini-1.5-flash'))
        elif config.get('openrouter_api_key'):
            success, response = generate_with_openrouter(prompt, config.get('openrouter_api_key'), config.get('selected_model', 'meta-llama/llama-3.2-3b-instruct:free'))
        else:
            return None
        
        if success and response:
            titles = [line.strip() for line in response.strip().split('\n') if line.strip()]
            # Ensure we have the right number of titles
            if len(titles) >= num_chapters:
                return titles[:num_chapters]
            else:
                # Pad with generic titles if needed
                while len(titles) < num_chapters:
                    titles.append(f"Chapter {len(titles) + 1}")
                return titles
        
        return None
    except Exception as e:
        logging.error(f"Error generating chapter titles: {e}")
        return None

def start_ai_content_generation(project_id, config):
    """Start AI content generation for a project in background"""
    import threading
    
    def generate_content():
        try:
            project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
            if not os.path.exists(project_file):
                return
            
            with open(project_file, 'r') as f:
                project = json.load(f)
            
            project['generation_status'] = 'generating'
            
            # Save status update
            with open(project_file, 'w') as f:
                json.dump(project, f, indent=2)
            
            # Generate content for each chapter
            for i, chapter in enumerate(project['chapters']):
                if chapter['status'] == 'pending':
                    chapter_prompt = f"""Write compelling content for this chapter:

Book Title: {project.get('title', project.get('name', 'Untitled'))}
Book Topic: {project['topic']}
Chapter Title: {chapter['title']}
Writing Style: {project['style']}
Language: {project['language']}

Please write detailed, engaging content for this chapter that:
- Fits the overall book theme
- Matches the {project['style']} writing style
- Is written in {project['language']}
- Is substantial (800-1200 words)
- Flows naturally from the chapter title
- Provides value to readers

Write the complete chapter content:"""

                    try:
                        # Generate content with AI
                        if config.get('ai_provider') == 'gemini' and config.get('gemini_api_key'):
                            success, content = generate_with_gemini(chapter_prompt, config.get('gemini_api_key'), config.get('gemini_model', 'gemini-1.5-flash'))
                        elif config.get('openrouter_api_key'):
                            success, content = generate_with_openrouter(chapter_prompt, config.get('openrouter_api_key'), config.get('selected_model', 'meta-llama/llama-3.2-3b-instruct:free'))
                        else:
                            success = False
                            content = f"AI content generation unavailable. Please edit this chapter manually.\n\nChapter: {chapter['title']}\n\nAdd your content here..."
                        
                        if success and content:
                            chapter['content'] = content
                            chapter['status'] = 'completed'
                            chapter['word_count'] = len(content.split())
                        else:
                            chapter['content'] = f"Content generation failed for {chapter['title']}. Please edit manually."
                            chapter['status'] = 'failed'
                    except Exception as e:
                        logging.error(f"Error generating content for chapter {i+1}: {e}")
                        chapter['content'] = f"Content generation failed for {chapter['title']}. Please edit manually."
                        chapter['status'] = 'failed'
                
                # Save progress after each chapter
                with open(project_file, 'w') as f:
                    json.dump(project, f, indent=2)
            
            # Mark project as completed
            project['generation_status'] = 'completed'
            project['last_modified'] = datetime.now().isoformat()
            
            with open(project_file, 'w') as f:
                json.dump(project, f, indent=2)
                
        except Exception as e:
            logging.error(f"Background content generation failed: {e}")
    
    # Start generation in background thread
    thread = threading.Thread(target=generate_content)
    thread.daemon = True
    thread.start()

def save_config(config):
    """Save configuration to config.json"""
    with open('config.json', 'w') as f:
        json.dump(config, f, indent=2)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def verify_license(license_key, email, machine_id):
    """Verify license with the API"""
    try:
        headers = {
            "x-vercel-protection-bypass": PROTECTION_BYPASS_TOKEN,
            "Content-Type": "application/json"
        }
        data = {
            "license_key": license_key,
            "email": email,
            "machine_id": machine_id
        }
        response = requests.post(API_URL, json=data, headers=headers, timeout=10)
        res_json = response.json()

        if response.status_code == 200 and res_json.get("success"):
            product_name = res_json.get("product_name", "")
            if product_name.lower() == REQUIRED_PRODUCT:
                return True, res_json
            else:
                return False, {"error": f"This license is not for {REQUIRED_PRODUCT}"}
        else:
            return False, res_json
    except Exception as e:
        return False, {"error": str(e)}

def generate_with_openrouter(prompt, api_key, model="openai/gpt-3.5-turbo"):
    """Generate content using OpenRouter API with retry logic and fallback models"""
    import time
    
    # List of free models to try as fallbacks
    free_models = [
        model,  # Try the requested model first
        "microsoft/phi-3-mini-128k-instruct:free",
        "huggingface/CodeLlama-7b-Instruct-hf:free", 
        "openchat/openchat-7b:free",
        "openai/gpt-3.5-turbo"  # Fallback to paid model if available
    ]
    
    for attempt, current_model in enumerate(free_models):
        try:
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }
            data = {
                "model": current_model,
                "messages": [
                    {"role": "user", "content": prompt}
                ]
            }
            
            response = requests.post("https://openrouter.ai/api/v1/chat/completions", 
                                   json=data, headers=headers, timeout=45)
            
            if response.status_code == 200:
                result = response.json()
                if 'choices' in result and len(result['choices']) > 0:
                    return True, result['choices'][0]['message']['content']
                else:
                    return False, f"No content in API response: {result}"
            elif response.status_code == 429:
                # Rate limited, try next model
                logging.warning(f"Rate limited on model {current_model}, trying next...")
                if attempt < len(free_models) - 1:
                    time.sleep(2)  # Brief wait before trying next model
                    continue
                else:
                    return False, f"All models rate limited. Please wait a few minutes and try again."
            else:
                # Other HTTP error, try next model
                logging.warning(f"HTTP {response.status_code} on model {current_model}: {response.text}")
                if attempt < len(free_models) - 1:
                    continue
                else:
                    return False, f"API Error: {response.status_code} - {response.text}"
                    
        except Exception as e:
            logging.error(f"Exception with model {current_model}: {str(e)}")
            if attempt < len(free_models) - 1:
                continue
            else:
                return False, f"Connection error: {str(e)}"
    
    return False, "All retry attempts failed"

def generate_with_gemini(prompt, api_key, model="gemini-2.5-flash"):
    """Generate content using Google Gemini API with the new google-genai library"""
    try:
        import os
        from google import genai
        from google.genai import types
        
        # Use API key from settings or environment
        if not api_key:
            api_key = os.environ.get("GEMINI_API_KEY")
        
        if not api_key:
            return False, "Gemini API key not configured. Please add it in Settings."
        
        client = genai.Client(api_key=api_key)
        
        response = client.models.generate_content(
            model=model,
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=0.7,
                max_output_tokens=4096,
            )
        )
        
        if response.text:
            return True, response.text
        else:
            return False, "No content generated by Gemini"
            
    except ImportError:
        return False, "Google Gemini library not installed. Please install google-genai."
    except Exception as e:
        logging.error(f"Gemini API exception: {str(e)}")
        return False, f"Gemini connection error: {str(e)}"

def generate_content(prompt, config):
    """Generate content using the selected AI provider"""
    ai_provider = config.get('ai_provider', 'openrouter')
    
    if ai_provider == 'gemini':
        api_key = config.get('gemini_api_key', '')
        model = config.get('gemini_model', 'gemini-1.5-flash')
        if not api_key:
            return False, "Gemini API key not configured"
        return generate_with_gemini(prompt, api_key, model)
    else:
        # Default to OpenRouter
        api_key = config.get('openrouter_api_key', '')
        model = config.get('selected_model', 'meta-llama/llama-3.2-3b-instruct:free')
        if not api_key:
            return False, "OpenRouter API key not configured"
        return generate_with_openrouter(prompt, api_key, model)

@app.route('/')
def index():
    config = load_config()
    
    # Check if license is activated
    if not config.get('license_activated', False):
        return render_template('index.html', show_license=True, config=config)
    
    # Get recent projects
    recent_projects = []
    try:
        for filename in os.listdir(PROJECTS_FOLDER):
            if filename.endswith('.json'):
                filepath = os.path.join(PROJECTS_FOLDER, filename)
                with open(filepath, 'r') as f:
                    project = json.load(f)
                    project['filename'] = filename
                    recent_projects.append(project)
        recent_projects.sort(key=lambda x: x.get('last_modified', ''), reverse=True)
        recent_projects = recent_projects[:5]  # Show only 5 recent projects
    except Exception as e:
        logging.error(f"Error loading recent projects: {e}")
    
    return render_template('index.html', config=config, recent_projects=recent_projects)

@app.route('/mood-tracker')
def mood_tracker():
    """Writing Mood Tracker page"""
    config = load_config()
    
    # Check if license is activated
    if not config.get('license_activated', False):
        return redirect(url_for('index'))
    
    return render_template('mood_tracker.html', config=config)

@app.route('/activate_license', methods=['POST'])
def activate_license():
    license_key = request.form.get('license_key', '').strip()
    email = request.form.get('email', '').strip()
    machine_id = get_machine_id()
    
    if not license_key or not email:
        flash('Please fill in all fields', 'error')
        return redirect(url_for('index'))
    
    success, result = verify_license(license_key, email, machine_id)
    
    if success:
        config = load_config()
        config.update({
            'license_activated': True,
            'license_key': license_key,
            'email': email,
            'license_info': result
        })
        save_config(config)
        flash(f'License activated successfully! License Type: {result.get("license_type", "Unknown")}', 'success')
    else:
        flash(f'License activation failed: {result.get("error", "Unknown error")}', 'error')
    
    return redirect(url_for('index'))

@app.route('/settings')
def settings():
    config = load_config()
    if not config.get('license_activated', False):
        flash('Please activate your license first', 'error')
        return redirect(url_for('index'))
    
    return render_template('settings.html', config=config)

@app.route('/save_settings', methods=['POST'])
def save_settings():
    config = load_config()
    config['ai_provider'] = request.form.get('ai_provider', 'openrouter')
    config['openrouter_api_key'] = request.form.get('openrouter_api_key', '')
    config['gemini_api_key'] = request.form.get('gemini_api_key', '')
    config['selected_model'] = request.form.get('model', 'meta-llama/llama-3.2-3b-instruct:free')
    config['gemini_model'] = request.form.get('gemini_model', 'gemini-1.5-flash')
    save_config(config)
    flash('Settings saved successfully! AI provider updated.', 'success')
    return redirect(url_for('settings'))

@app.route('/save_author_bio', methods=['POST'])
def save_author_bio():
    config = load_config()
    config['default_author_bio'] = request.form.get('default_author_bio', '')
    config['use_default_bio'] = request.form.get('use_default_bio') == 'true'
    save_config(config)
    flash('Default author bio saved successfully!', 'success')
    return redirect(url_for('settings'))

@app.route('/check_ai_provider_status')
def check_ai_provider_status():
    """Check the current AI provider configuration status"""
    config = load_config()
    ai_provider = config.get('ai_provider', 'openrouter')
    
    if ai_provider == 'gemini':
        api_key = config.get('gemini_api_key', '')
        model = config.get('gemini_model', 'gemini-1.5-flash')
        configured = bool(api_key)
        provider_name = 'Google Gemini'
    else:
        api_key = config.get('openrouter_api_key', '')
        model = config.get('selected_model', 'meta-llama/llama-3.2-3b-instruct:free')
        configured = bool(api_key)
        provider_name = 'OpenRouter'
    
    return jsonify({
        'provider': ai_provider,
        'provider_name': provider_name,
        'model': model,
        'configured': configured,
        'status': 'ready' if configured else 'not_configured'
    })

@app.route('/create_manual_book', methods=['POST'])
def create_manual_book():
    """Create a manual book project with enhanced options"""
    config = load_config()
    if not config.get('license_activated', False):
        flash('Please activate your license first', 'error')
        return redirect(url_for('index'))
    
    try:
        # Get form data
        book_title = request.form.get('book_title', '').strip()
        topic = request.form.get('topic', '').strip()
        author_bio = request.form.get('author_bio', '').strip()
        language = request.form.get('language', 'English')
        num_chapters = int(request.form.get('chapters', 8))
        style = request.form.get('style', 'professional')
        generation_mood = request.form.get('generation_mood', '')
        
        # Get generation options
        chapter_titles_method = request.form.get('chapter_titles_method', 'manual')
        content_method = request.form.get('content_method', 'manual')
        
        if not book_title or not topic:
            flash('Please provide both book title and topic', 'error')
            return redirect(url_for('index'))
        
        # Create project
        project_id = str(uuid.uuid4())[:8]
        project = {
            'id': project_id,
            'title': book_title,
            'name': book_title,
            'topic': topic,
            'author_bio': author_bio,
            'language': language,
            'style': style,
            'num_chapters': num_chapters,
            'creation_method': 'manual',
            'chapter_titles_method': chapter_titles_method,
            'content_method': content_method,
            'generation_status': 'draft',
            'created_at': datetime.now().isoformat(),
            'last_modified': datetime.now().isoformat(),
            'chapters': [],
            'generation_mood': generation_mood,
            'cover_image': None
        }
        
        # Handle chapter titles based on method
        if chapter_titles_method == 'ai':
            # Generate AI chapter titles
            try:
                chapter_titles = generate_chapter_titles_ai(book_title, topic, language, num_chapters, style, config)
                if chapter_titles:
                    for i, title in enumerate(chapter_titles, 1):
                        chapter = {
                            'id': str(uuid.uuid4())[:8],
                            'number': i,
                            'title': title,
                            'content': f'Content for {title} - Click to edit and add your own content here.' if content_method == 'manual' else '',
                            'status': 'draft' if content_method == 'manual' else 'pending',
                            'word_count': 0
                        }
                        project['chapters'].append(chapter)
                else:
                    # Fallback to manual if AI fails
                    for i in range(1, num_chapters + 1):
                        chapter = {
                            'id': str(uuid.uuid4())[:8],
                            'number': i,
                            'title': f'Chapter {i}',
                            'content': f'Content for Chapter {i} - Click to edit and add your own content here.',
                            'status': 'draft',
                            'word_count': 0
                        }
                        project['chapters'].append(chapter)
                    flash('AI chapter title generation failed, created manual chapters', 'warning')
            except Exception as e:
                logging.error(f"AI chapter title generation failed: {e}")
                # Create manual chapters as fallback
                for i in range(1, num_chapters + 1):
                    chapter = {
                        'id': str(uuid.uuid4())[:8],
                        'number': i,
                        'title': f'Chapter {i}',
                        'content': f'Content for Chapter {i} - Click to edit and add your own content here.',
                        'status': 'draft',
                        'word_count': 0
                    }
                    project['chapters'].append(chapter)
                flash('AI unavailable, created manual chapters', 'warning')
        else:
            # Create manual chapters
            for i in range(1, num_chapters + 1):
                chapter = {
                    'id': str(uuid.uuid4())[:8],
                    'number': i,
                    'title': f'Chapter {i}',
                    'content': f'Content for Chapter {i} - Click to edit and add your own content here.',
                    'status': 'draft',
                    'word_count': 0
                }
                project['chapters'].append(chapter)
        
        # If AI content generation is requested, mark for background processing
        if content_method == 'ai' and project['chapters']:
            project['generation_status'] = 'pending'
            for chapter in project['chapters']:
                chapter['status'] = 'pending'
        
        # Save project
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
        
        # Start AI content generation in background if requested
        if content_method == 'ai':
            flash(f'Book project "{book_title}" created! AI content generation starting...', 'success')
            # Start background generation
            try:
                start_ai_content_generation(project_id, config)
            except Exception as e:
                logging.error(f"Failed to start AI content generation: {e}")
                flash('Project created but AI content generation failed to start', 'warning')
        else:
            flash(f'Book project "{book_title}" created successfully! Start editing chapters.', 'success')
        
        return redirect(url_for('project_view', project_id=project_id))
        
    except Exception as e:
        logging.error(f"Error creating manual project: {e}")
        flash(f'Error creating project: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/save_mood', methods=['POST'])
def save_mood():
    """Save user's writing mood for today"""
    try:
        data = request.get_json()
        mood = data.get('mood')
        note = data.get('note', '')
        
        if not mood:
            return jsonify({'success': False, 'message': 'Mood is required'})
        
        # Load or create mood data
        mood_file = 'mood_data.json'
        try:
            with open(mood_file, 'r') as f:
                mood_data = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            mood_data = {'moods': []}
        
        # Get today's date
        today = datetime.now().strftime('%Y-%m-%d')
        
        # Check if mood already exists for today
        existing_mood_index = None
        for i, entry in enumerate(mood_data['moods']):
            if entry.get('date') == today:
                existing_mood_index = i
                break
        
        # Create mood entry
        mood_entry = {
            'date': today,
            'mood': mood,
            'note': note,
            'timestamp': datetime.now().isoformat()
        }
        
        # Update or add mood entry
        if existing_mood_index is not None:
            mood_data['moods'][existing_mood_index] = mood_entry
        else:
            mood_data['moods'].append(mood_entry)
        
        # Keep only last 30 days
        mood_data['moods'] = mood_data['moods'][-30:]
        
        # Save mood data
        with open(mood_file, 'w') as f:
            json.dump(mood_data, f, indent=2)
        
        return jsonify({'success': True, 'message': 'Mood saved successfully'})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error saving mood: {str(e)}'})

@app.route('/get_mood_history')
def get_mood_history():
    """Get user's mood history"""
    try:
        mood_file = 'mood_data.json'
        try:
            with open(mood_file, 'r') as f:
                mood_data = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            mood_data = {'moods': []}
        
        # Sort by date (newest first)
        mood_data['moods'].sort(key=lambda x: x['date'], reverse=True)
        
        return jsonify({'success': True, 'moods': mood_data['moods']})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error loading mood history: {str(e)}'})

@app.route('/project/<project_id>/save_project_mood', methods=['POST'])
def save_project_mood(project_id):
    """Save user's writing mood for specific project"""
    try:
        data = request.get_json()
        mood = data.get('mood')
        note = data.get('note', '')
        
        if not mood:
            return jsonify({'success': False, 'message': 'Mood is required'})
        
        # Load or create project-specific mood data
        mood_file = f'project_{project_id}_mood.json'
        try:
            with open(mood_file, 'r') as f:
                mood_data = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            mood_data = {'project_id': project_id, 'moods': []}
        
        # Get today's date
        today = datetime.now().strftime('%Y-%m-%d')
        
        # Check if mood already exists for today
        existing_mood_index = None
        for i, entry in enumerate(mood_data['moods']):
            if entry.get('date') == today:
                existing_mood_index = i
                break
        
        # Create mood entry
        mood_entry = {
            'date': today,
            'mood': mood,
            'note': note,
            'timestamp': datetime.now().isoformat(),
            'project_id': project_id
        }
        
        # Update or add mood entry
        if existing_mood_index is not None:
            mood_data['moods'][existing_mood_index] = mood_entry
        else:
            mood_data['moods'].append(mood_entry)
        
        # Keep only last 30 days
        mood_data['moods'] = mood_data['moods'][-30:]
        
        # Save mood data
        with open(mood_file, 'w') as f:
            json.dump(mood_data, f, indent=2)
        
        return jsonify({'success': True, 'message': 'Project mood saved successfully'})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error saving project mood: {str(e)}'})

@app.route('/project/<project_id>/get_project_mood_history')
def get_project_mood_history(project_id):
    """Get project-specific mood history"""
    try:
        mood_file = f'project_{project_id}_mood.json'
        try:
            with open(mood_file, 'r') as f:
                mood_data = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            mood_data = {'project_id': project_id, 'moods': []}
        
        # Sort by date (newest first)
        mood_data['moods'].sort(key=lambda x: x['date'], reverse=True)
        
        return jsonify({'success': True, 'moods': mood_data['moods'], 'project_id': project_id})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error loading project mood history: {str(e)}'})

@app.route('/create_project', methods=['POST'])
def create_project():
    config = load_config()
    if not config.get('license_activated', False):
        flash('Please activate your license first', 'error')
        return redirect(url_for('index'))
    
    project_name = request.form.get('project_name', '').strip()
    book_topic = request.form.get('book_topic', '').strip()
    author_bio = request.form.get('author_bio', '').strip()
    language = request.form.get('language', 'English')
    num_chapters = int(request.form.get('num_chapters', 10))
    
    if not project_name or not book_topic:
        flash('Project name and book topic are required', 'error')
        return redirect(url_for('index'))
    
    # Handle cover image upload
    cover_image = None
    if 'cover_image' in request.files:
        file = request.files['cover_image']
        if file and file.filename and file.filename != '' and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filename = f"{uuid.uuid4()}_{filename}"
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            cover_image = filename
    
    # Create new project
    project_id = str(uuid.uuid4())
    project = {
        'id': project_id,
        'name': project_name,
        'topic': book_topic,
        'author_bio': author_bio,
        'language': language,
        'num_chapters': num_chapters,
        'cover_image': cover_image,
        'chapters': [],
        'created_at': datetime.now().isoformat(),
        'last_modified': datetime.now().isoformat(),
        'generation_status': 'pending'
    }
    
    # Save project
    project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
    with open(project_file, 'w') as f:
        json.dump(project, f, indent=2)
    
    return redirect(url_for('project_view', project_id=project_id))

def enhance_book_title(title, topic, language="English"):
    """Enhance book title using AI"""
    try:
        config = load_config()
        
        prompt = f"""You are a professional book title consultant. Enhance and improve this book title to make it more engaging, marketable, and compelling.

Original Title: "{title}"
Book Topic/Description: {topic}
Target Language: {language}

Guidelines:
- Make it catchy and marketable
- Keep it relevant to the topic
- Consider the target audience
- Add subtitle if beneficial
- Maximum 100 characters total
- Return ONLY the enhanced title, nothing else

Enhanced title:"""

        if config['ai_provider'] == 'gemini' and config.get('gemini_api_key'):
            import google.genai as genai
            client = genai.Client(api_key=config['gemini_api_key'])
            response = client.models.generate_content(
                model=config.get('gemini_model', 'gemini-2.5-flash'),
                contents=prompt
            )
            return response.text.strip() if response.text else title
        
        elif config.get('openrouter_api_key'):
            response = requests.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers={"Authorization": f"Bearer {config['openrouter_api_key']}"},
                json={
                    "model": config.get('selected_model', 'anthropic/claude-3.5-sonnet:beta'),
                    "messages": [{"role": "user", "content": prompt}],
                    "max_tokens": 200
                }
            )
            return response.json()['choices'][0]['message']['content'].strip()
        
        return title
        
    except Exception as e:
        logging.error(f"Error enhancing title: {e}")
        return title

def generate_book_description(title, topic, author_bio="", language="English"):
    """Generate professional book description using AI"""
    try:
        config = load_config()
        
        prompt = f"""You are a professional book marketing copywriter. Create a compelling, professional book description for Amazon KDP and other book platforms.

Book Title: {title}
Topic/Content: {topic}
Author Info: {author_bio if author_bio else "Not provided"}
Language: {language}

Create a description that includes:
1. A compelling hook that grabs attention
2. What readers will learn or gain
3. Target audience
4. Key benefits and outcomes
5. Professional tone suitable for book marketing
6. 150-300 words total

Write in {language}. Return ONLY the book description:"""

        if config['ai_provider'] == 'gemini' and config.get('gemini_api_key'):
            import google.genai as genai
            client = genai.Client(api_key=config['gemini_api_key'])
            response = client.models.generate_content(
                model=config.get('gemini_model', 'gemini-2.5-flash'),
                contents=prompt
            )
            return response.text.strip() if response.text else "Professional book covering essential topics and insights."
        
        elif config.get('openrouter_api_key'):
            response = requests.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers={"Authorization": f"Bearer {config['openrouter_api_key']}"},
                json={
                    "model": config.get('selected_model', 'anthropic/claude-3.5-sonnet:beta'),
                    "messages": [{"role": "user", "content": prompt}],
                    "max_tokens": 500
                }
            )
            return response.json()['choices'][0]['message']['content'].strip()
        
        return "Professional book covering essential topics and insights."
        
    except Exception as e:
        logging.error(f"Error generating description: {e}")
        return "Professional book covering essential topics and insights."

@app.route('/create_ai_book', methods=['POST'])
def create_ai_book():
    config = load_config()
    if not config.get('license_activated', False):
        flash('Please activate your license first', 'error')
        return redirect(url_for('index'))
    
    book_title = request.form.get('book_title', '').strip()
    topic = request.form.get('topic', '').strip()
    author_bio = request.form.get('author_bio', '').strip()
    language = request.form.get('language', 'English')
    chapters = int(request.form.get('chapters', 8))
    style = request.form.get('style', 'professional')
    action = request.form.get('action', 'generate_titles')
    generation_mood = request.form.get('generation_mood', '')
    
    if not book_title or not topic:
        flash('Book title and topic are required', 'error')
        return redirect(url_for('index'))
    
    # Enhance the book title using AI
    enhanced_title = enhance_book_title(book_title, topic, language)
    
    # Generate book description using AI
    book_description = generate_book_description(enhanced_title, topic, author_bio, language)
    
    # Use enhanced title as project name
    project_name = enhanced_title[:80]
    
    # Create new project
    project_id = str(uuid.uuid4())
    project = {
        'id': project_id,
        'name': project_name,
        'title': enhanced_title,
        'original_title': book_title,
        'description': book_description,
        'topic': topic,
        'author_bio': author_bio,
        'language': language,
        'num_chapters': chapters,
        'writing_style': style,
        'cover_image': None,
        'chapters': [],
        'created_at': datetime.now().isoformat(),
        'last_modified': datetime.now().isoformat(),
        'generation_status': 'pending',
        'ai_generated': True,
        'generation_action': action,
        'generation_mood': generation_mood
    }
    
    # Save project
    project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
    with open(project_file, 'w') as f:
        json.dump(project, f, indent=2)
    
    # Start generation based on action
    if action in ['generate_titles', 'generate_full']:
        return redirect(url_for('generate_chapters', project_id=project_id))
    
    return redirect(url_for('project_view', project_id=project_id))

@app.route('/project/<project_id>')
def project_view(project_id):
    config = load_config()
    if not config.get('license_activated', False):
        flash('Please activate your license first', 'error')
        return redirect(url_for('index'))
    
    try:
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        return render_template('project.html', project=project, config=config)
    except FileNotFoundError:
        flash('Project not found', 'error')
        return redirect(url_for('index'))

@app.route('/project/<project_id>/upload_cover', methods=['POST'])
def upload_project_cover(project_id):
    """Upload or update cover image for existing project"""
    config = load_config()
    if not config.get('license_activated', False):
        flash('Please activate your license first', 'error')
        return redirect(url_for('index'))
    
    try:
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        if 'cover_image' not in request.files:
            flash('No file selected', 'error')
            return redirect(url_for('project_view', project_id=project_id))
        
        file = request.files['cover_image']
        if file.filename == '':
            flash('No file selected', 'error')
            return redirect(url_for('project_view', project_id=project_id))
        
        if file and file.filename and allowed_file(file.filename):
            # Remove old cover image if exists
            if project.get('cover_image'):
                old_cover_path = os.path.join(UPLOAD_FOLDER, project['cover_image'])
                if os.path.exists(old_cover_path):
                    os.remove(old_cover_path)
            
            # Save new cover image
            filename = secure_filename(file.filename)
            filename = f"{project_id}_{filename}"
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            file.save(file_path)
            
            # Resize image if needed
            try:
                with Image.open(file_path) as img:
                    # Convert to RGB if necessary
                    if img.mode in ('RGBA', 'LA', 'P'):
                        img = img.convert('RGB')
                    
                    # Resize to reasonable dimensions while maintaining aspect ratio
                    img.thumbnail((800, 1200), Image.Resampling.LANCZOS)
                    img.save(file_path, 'JPEG', quality=85)
                    
            except Exception as e:
                flash(f'Error processing image: {str(e)}', 'error')
                return redirect(url_for('project_view', project_id=project_id))
            
            # Update project data
            project['cover_image'] = filename
            project['last_modified'] = datetime.now().isoformat()
            
            with open(project_file, 'w') as f:
                json.dump(project, f, indent=2)
            
            flash('Cover image updated successfully!', 'success')
        else:
            flash('Invalid file type. Please upload an image file.', 'error')
        
        return redirect(url_for('project_view', project_id=project_id))
        
    except Exception as e:
        flash(f'Error updating cover image: {str(e)}', 'error')
        return redirect(url_for('project_view', project_id=project_id))

@app.route('/delete_project/<project_id>', methods=['POST', 'DELETE'])
def delete_project(project_id):
    """Delete a project and all its associated files"""
    config = load_config()
    if not config.get('license_activated', False):
        flash('Please activate your license first', 'error')
        return redirect(url_for('index'))
    
    try:
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        
        # Load project to get cover image path
        if os.path.exists(project_file):
            with open(project_file, 'r') as f:
                project = json.load(f)
            
            # Remove cover image if it exists
            if project.get('cover_image'):
                cover_path = os.path.join(UPLOAD_FOLDER, project['cover_image'])
                if os.path.exists(cover_path):
                    os.remove(cover_path)
            
            # Remove project file
            os.remove(project_file)
            
            flash('Project deleted successfully!', 'success')
        else:
            flash('Project not found', 'error')
            
    except Exception as e:
        flash(f'Error deleting project: {str(e)}', 'error')
    
    return redirect(url_for('index'))

@app.route('/project/<project_id>/preview')
def book_preview(project_id):
    """Live book preview with real-time editing"""
    config = load_config()
    if not config.get('license_activated', False):
        flash('Please activate your license first', 'error')
        return redirect(url_for('index'))
    
    try:
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        return render_template('book_preview.html', project=project, config=config)
    except FileNotFoundError:
        flash('Project not found', 'error')
        return redirect(url_for('index'))

@app.route('/api/update_project/<project_id>', methods=['POST'])
def update_project_content(project_id):
    """Update project content via AJAX"""
    config = load_config()
    if not config.get('license_activated', False):
        return jsonify({'success': False, 'message': 'License not activated'})
    
    try:
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        data = request.get_json()
        
        # Update title if provided
        if 'title' in data:
            project['name'] = data['title']
        
        # Update topic/description if provided
        if 'topic' in data:
            project['topic'] = data['topic']
        
        # Handle form data for legacy support
        edit_type = request.form.get('editType') if request.form else None
        
        if edit_type == 'cover':
            project['name'] = request.form.get('name', project['name'])
            project['topic'] = request.form.get('topic', project['topic'])
        elif edit_type == 'chapter':
            chapter_index_str = request.form.get('chapterIndex')
            if chapter_index_str is not None:
                chapter_index = int(chapter_index_str)
                if 0 <= chapter_index < len(project.get('chapters', [])):
                    project['chapters'][chapter_index]['title'] = request.form.get('title')
                    project['chapters'][chapter_index]['content'] = request.form.get('content')
        
        project['last_modified'] = datetime.now().isoformat()
        
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
        
        return jsonify({'success': True, 'message': 'Project updated successfully'})
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/update_chapter', methods=['POST'])
def update_chapter_content():
    """Update specific chapter content"""
    config = load_config()
    if not config.get('license_activated', False):
        return jsonify({'success': False, 'message': 'License not activated'})
    
    try:
        data = request.get_json()
        project_id = data.get('project_id')
        chapter_id = data.get('chapter_id')
        content = data.get('content')
        
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        # Find and update the chapter
        for chapter in project.get('chapters', []):
            if chapter.get('id') == chapter_id:
                chapter['content'] = content
                break
        
        project['last_modified'] = datetime.now().isoformat()
        
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
        
        return jsonify({'success': True, 'message': 'Chapter updated successfully'})
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/generate_chapters/<project_id>')
def generate_chapters(project_id):
    config = load_config()
    ai_provider = config.get('ai_provider', 'openrouter')
    
    # Check if the selected AI provider has API key configured
    if ai_provider == 'gemini':
        api_key = config.get('gemini_api_key', '')
        if not api_key:
            flash('Please configure your Gemini API key in settings', 'error')
            return redirect(url_for('project_view', project_id=project_id))
    else:
        api_key = config.get('openrouter_api_key', '')
        if not api_key:
            flash('Please configure your OpenRouter API key in settings', 'error')
            return redirect(url_for('project_view', project_id=project_id))
    
    try:
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        # Start background generation
        thread = threading.Thread(target=generate_chapters_background, 
                                args=(project_id, project, config))
        thread.daemon = True
        thread.start()
        
        flash('Chapter generation started! Please wait...', 'info')
        return redirect(url_for('project_view', project_id=project_id))
        
    except Exception as e:
        flash(f'Error starting generation: {str(e)}', 'error')
        return redirect(url_for('project_view', project_id=project_id))

def generate_chapters_background(project_id, project, config):
    """Background task to generate chapters"""
    project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
    try:
        # Update status
        project['generation_status'] = 'generating_titles'
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
        
        # Enhanced description step before generating titles
        project['generation_status'] = 'enhancing_description'
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
        
        # First enhance the topic description for better context
        description_prompt = f"""Analyze and enhance this book topic: "{project['topic']}" in {project['language']}.
        
        Create a comprehensive book description that includes:
        1. Main theme and purpose
        2. Target audience 
        3. Key concepts to be covered
        4. Writing style and tone
        5. Learning outcomes or takeaways
        
        Return a detailed, professional description that will guide chapter creation."""
        
        success, enhanced_description = generate_content(description_prompt, config)
        
        if success:
            project['enhanced_description'] = enhanced_description.strip()
        else:
            project['enhanced_description'] = project['topic']  # Fallback to original topic
        
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
        
        # Update status for title generation
        project['generation_status'] = 'generating_titles'
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
        
        # Generate chapter titles using enhanced description
        titles_prompt = f"""Based on this enhanced book description: "{project.get('enhanced_description', project['topic'])}"
        
        Generate {project['num_chapters']} compelling, specific chapter titles in {project['language']} that:
        1. Follow a logical progression
        2. Cover all key aspects mentioned in the description
        3. Are engaging and professional
        4. Build upon each other naturally
        5. Appeal to the target audience
        
        Return only the titles, one per line, numbered from 1 to {project['num_chapters']}."""
        
        success, titles_result = generate_content(titles_prompt, config)
        
        if not success:
            project['generation_status'] = f'error: {titles_result}'
            with open(project_file, 'w') as f:
                json.dump(project, f, indent=2)
            return
        
        # Parse chapter titles
        titles = []
        for line in titles_result.strip().split('\n'):
            line = line.strip()
            if line and ('.' in line or ':' in line):
                # Remove numbering
                title = line.split('.', 1)[-1].split(':', 1)[-1].strip()
                if title:
                    titles.append(title)
        
        if len(titles) < project['num_chapters']:
            # Fill missing titles
            for i in range(len(titles), project['num_chapters']):
                titles.append(f"Chapter {i+1}")
        
        # Initialize chapters
        project['chapters'] = []
        for i, title in enumerate(titles[:project['num_chapters']]):
            project['chapters'].append({
                'id': str(uuid.uuid4()),
                'number': i + 1,
                'title': title,
                'content': '',
                'status': 'pending'
            })
        
        project['generation_status'] = 'generating_content'
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
        
        # Generate content for each chapter
        for i, chapter in enumerate(project['chapters']):
            chapter['status'] = 'generating'
            with open(project_file, 'w') as f:
                json.dump(project, f, indent=2)
            
            content_prompt = f"""Write comprehensive, professional content for Chapter {chapter['number']}: "{chapter['title']}" 
            
            Book Context: {project.get('enhanced_description', project['topic'])}
            Language: {project['language']} 

            IMPORTANT FORMATTING REQUIREMENTS:
            - Write in clean, professional prose suitable for Amazon KDP publishing
            - Do NOT use any Markdown formatting (no ##, ###, **, *, _, etc.)
            - Use proper paragraph breaks for readability
            - Write section headings as regular text with proper capitalization
            - The content should be detailed, engaging, and approximately 1000-1500 words
            - Use professional book formatting with clear paragraph structure
            - No bullet points with asterisks - use proper paragraph flow instead
            
            Create content that flows naturally like a professionally published book."""
            
            success, content_result = generate_content(content_prompt, config)
            
            if success:
                # Clean markdown formatting from generated content
                import re
                clean_content = content_result.strip()
                # Remove markdown headers
                clean_content = re.sub(r'^#{1,6}\s+', '', clean_content, flags=re.MULTILINE)
                # Remove bold/italic markers
                clean_content = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', clean_content)
                clean_content = re.sub(r'_{1,2}([^_]+)_{1,2}', r'\1', clean_content)
                # Remove bullet point markers
                clean_content = re.sub(r'^\*\s+', '', clean_content, flags=re.MULTILINE)
                clean_content = re.sub(r'^\-\s+', '', clean_content, flags=re.MULTILINE)
                
                chapter['content'] = clean_content
                chapter['status'] = 'completed'
            else:
                chapter['content'] = f"Error generating content: {content_result}"
                chapter['status'] = 'error'
            
            with open(project_file, 'w') as f:
                json.dump(project, f, indent=2)
        
        # Mark as completed
        project['generation_status'] = 'completed'
        project['last_modified'] = datetime.now().isoformat()
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
            
    except Exception as e:
        project['generation_status'] = f'error: {str(e)}'
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)

@app.route('/api/projects')
def api_projects():
    """Get all projects with enhanced metadata for homepage"""
    try:
        projects = []
        if os.path.exists(PROJECTS_FOLDER):
            for filename in os.listdir(PROJECTS_FOLDER):
                if filename.endswith('.json'):
                    project_file = os.path.join(PROJECTS_FOLDER, filename)
                    try:
                        with open(project_file, 'r') as f:
                            project = json.load(f)
                        
                        # Enhanced metadata
                        project['filename'] = filename
                        project['id'] = filename[:-5]  # Remove .json extension
                        
                        # Add file timestamps if missing
                        if 'created_date' not in project:
                            project['created_date'] = datetime.fromtimestamp(os.path.getctime(project_file)).isoformat()
                        
                        if 'last_modified' not in project:
                            project['last_modified'] = datetime.fromtimestamp(os.path.getmtime(project_file)).isoformat()
                        
                        # Calculate completion status
                        if project.get('generation_status') == 'completed':
                            project['status'] = 'completed'
                        elif project.get('chapters') and len(project.get('chapters', [])) > 0:
                            completed_chapters = len([c for c in project.get('chapters', []) if c.get('content', '').strip()])
                            total_chapters = len(project.get('chapters', []))
                            project['completion_percentage'] = (completed_chapters / total_chapters * 100) if total_chapters > 0 else 0
                            project['status'] = 'completed' if completed_chapters == total_chapters else 'in_progress'
                        else:
                            project['status'] = 'draft'
                        
                        projects.append(project)
                    except Exception as e:
                        logging.warning(f"Error reading project file {filename}: {e}")
                        continue
        
        # Sort by last modified date (newest first)
        projects.sort(key=lambda x: x.get('last_modified', ''), reverse=True)
        
        return jsonify({'projects': projects})
    except Exception as e:
        logging.error(f"Error fetching projects: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/project_status/<project_id>')
def project_status(project_id):
    """Get project generation status via AJAX"""
    try:
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        completed_chapters = len([c for c in project.get('chapters', []) if c.get('status') == 'completed'])
        total_chapters = len(project.get('chapters', []))
        
        return jsonify({
            'status': project.get('generation_status', 'pending'),
            'chapters': project.get('chapters', []),
            'progress': (completed_chapters / max(total_chapters, 1) * 100) if total_chapters > 0 else 0,
            'completed_chapters': completed_chapters,
            'total_chapters': total_chapters
        })
    except FileNotFoundError:
        return jsonify({'status': 'error', 'message': 'Project not found'}), 404
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/regenerate_chapter/<project_id>/<chapter_id>')
def regenerate_chapter(project_id, chapter_id):
    """Regenerate a specific chapter"""
    config = load_config()
    ai_provider = config.get('ai_provider', 'openrouter')
    
    # Check if the selected AI provider has API key configured
    if ai_provider == 'gemini':
        api_key = config.get('gemini_api_key', '')
        if not api_key:
            return jsonify({'error': 'Please configure your Gemini API key in settings'}), 400
    else:
        api_key = config.get('openrouter_api_key', '')
        if not api_key:
            return jsonify({'error': 'Please configure your OpenRouter API key in settings'}), 400
    
    try:
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        # Find the chapter
        chapter = None
        for c in project['chapters']:
            if c['id'] == chapter_id:
                chapter = c
                break
        
        if not chapter:
            return jsonify({'error': 'Chapter not found'}), 404
        
        # Start background regeneration
        thread = threading.Thread(target=regenerate_single_chapter, 
                                args=(project_id, chapter_id, project, config))
        thread.daemon = True
        thread.start()
        
        return jsonify({'success': True, 'message': f'Chapter "{chapter["title"]}" regeneration started!'})
        
    except Exception as e:
        return jsonify({'error': f'Error starting regeneration: {str(e)}'}), 500

@app.route('/api/chapter_status/<project_id>/<chapter_id>')
def chapter_status(project_id, chapter_id):
    """Get status of a specific chapter"""
    try:
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        # Find the chapter
        chapter = None
        for c in project['chapters']:
            if c['id'] == chapter_id:
                chapter = c
                break
        
        if not chapter:
            return jsonify({'error': 'Chapter not found'}), 404
        
        return jsonify({
            'status': chapter.get('status', 'pending'),
            'title': chapter.get('title', ''),
            'content_length': len(chapter.get('content', ''))
        })
        
    except FileNotFoundError:
        return jsonify({'error': 'Project not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def regenerate_single_chapter(project_id, chapter_id, project, config):
    """Background task to regenerate a single chapter"""
    project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
    
    try:
        # Find and update the chapter
        for chapter in project['chapters']:
            if chapter['id'] == chapter_id:
                chapter['status'] = 'generating'
                with open(project_file, 'w') as f:
                    json.dump(project, f, indent=2)
                
                content_prompt = f"""Write comprehensive content for Chapter {chapter['number']}: "{chapter['title']}" 
                for a book about "{project['topic']}" in {project['language']}. 
                The content should be detailed, engaging, and approximately 1000-1500 words. 
                Use proper formatting with paragraphs and sections where appropriate."""
                
                success, content_result = generate_content(content_prompt, config)
                
                if success:
                    chapter['content'] = content_result.strip()
                    chapter['status'] = 'completed'
                else:
                    chapter['content'] = f"Error generating content: {content_result}"
                    chapter['status'] = 'error'
                
                project['last_modified'] = datetime.now().isoformat()
                with open(project_file, 'w') as f:
                    json.dump(project, f, indent=2)
                break
                
    except Exception as e:
        # Update chapter with error
        for chapter in project['chapters']:
            if chapter['id'] == chapter_id:
                chapter['status'] = 'error'
                chapter['content'] = f"Error: {str(e)}"
                with open(project_file, 'w') as f:
                    json.dump(project, f, indent=2)
                break

@app.route('/edit_chapter/<project_id>/<chapter_id>', methods=['POST'])
def edit_chapter(project_id, chapter_id):
    try:
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        new_title = request.form.get('title', '').strip()
        new_content = request.form.get('content', '').strip()
        
        # Find and update chapter
        for chapter in project['chapters']:
            if chapter['id'] == chapter_id:
                chapter['title'] = new_title
                chapter['content'] = new_content
                break
        
        project['last_modified'] = datetime.now().isoformat()
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
        
        flash('Chapter updated successfully!', 'success')
        return redirect(url_for('project_view', project_id=project_id))
        
    except Exception as e:
        flash(f'Error updating chapter: {str(e)}', 'error')
        return redirect(url_for('project_view', project_id=project_id))

def clean_chapter_content(content):
    """Clean and format chapter content for export"""
    if not content:
        return []
    
    lines = content.split('\n')
    cleaned_lines = []
    
    for line in lines:
        line = line.strip()
        # Skip empty lines, markdown headers, and duplicate titles
        if (line and 
            not line.startswith('#') and 
            not line.startswith('Chapter') and 
            not line.startswith('CHAPTER') and 
            not line.startswith('hapitre') and
            not (len(line) > 10 and line.isupper())):  # Skip all-caps lines
            
            # Clean markdown formatting
            line = line.replace('**', '').replace('*', '').replace('_', '').replace('##', '').replace('###', '')
            line = line.strip()
            if line and len(line) > 10:  # Only keep substantial content
                cleaned_lines.append(line)
    
    # Join all text and create proper paragraphs
    full_text = ' '.join(cleaned_lines)
    
    # Split into sentences roughly
    sentences = []
    current_sentence = ""
    
    for char in full_text:
        current_sentence += char
        if char in '.!?' and len(current_sentence.strip()) > 15:
            sentences.append(current_sentence.strip())
            current_sentence = ""
    
    if current_sentence.strip():
        sentences.append(current_sentence.strip())
    
    # Group sentences into paragraphs (3-4 sentences per paragraph)
    paragraphs = []
    current_paragraph = []
    
    for sentence in sentences:
        if sentence and len(sentence) > 10:
            current_paragraph.append(sentence)
            if len(current_paragraph) >= 3:
                paragraphs.append(' '.join(current_paragraph))
                current_paragraph = []
    
    if current_paragraph:
        paragraphs.append(' '.join(current_paragraph))
    
    return paragraphs

@app.route('/export_pdf/<project_id>')
def export_pdf(project_id):
    try:
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        # Clean chapter content for better formatting
        if project.get('chapters'):
            for chapter in project['chapters']:
                chapter['cleaned_paragraphs'] = clean_chapter_content(chapter.get('content', ''))
        
        # Generate HTML content
        html_content = render_template('book_export.html', project=project)
        
        # Fix relative URLs to absolute paths for PDF generation
        current_dir = os.path.dirname(os.path.abspath(__file__))
        static_path = os.path.join(current_dir, 'static', 'uploads')
        html_content = html_content.replace('/static/uploads/', f'file://{static_path}/')
        
        # Convert to PDF
        pdf_filename = f"{project['name']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        pdf_path = os.path.join(EXPORTS_FOLDER, pdf_filename)
        
        weasyprint.HTML(string=html_content, base_url=request.url_root).write_pdf(pdf_path)
        
        return send_file(pdf_path, as_attachment=True, download_name=pdf_filename)
        
    except Exception as e:
        flash(f'Error exporting PDF: {str(e)}', 'error')
        return redirect(url_for('project_view', project_id=project_id))

@app.route('/export_docx/<project_id>')
def export_docx(project_id):
    """Export project as DOCX file"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
        from docx.enum.style import WD_STYLE_TYPE
        
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        # Create new Word document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Cover page with image if available
        if project.get('cover_image'):
            cover_image_path = os.path.join(UPLOAD_FOLDER, project['cover_image'])
            if os.path.exists(cover_image_path):
                # Add cover image
                cover_paragraph = doc.add_paragraph()
                cover_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cover_run = cover_paragraph.add_run()
                try:
                    # Add image with full width
                    cover_run.add_picture(cover_image_path, width=Inches(6))
                except Exception as e:
                    # If image fails, fall back to text cover
                    pass
        
        # Title overlay or regular title
        doc.add_paragraph()  # Space
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(project['name'])
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        
        # Subtitle/topic
        if project.get('topic'):
            topic_paragraph = doc.add_paragraph()
            topic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            topic_run = topic_paragraph.add_run(project['topic'])
            topic_run.font.size = Pt(16)
            topic_run.font.italic = True
        
        doc.add_paragraph()  # Empty line
        
        # Book metadata
        meta_paragraph = doc.add_paragraph()
        meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_paragraph.add_run(f"Generated with BookGenPro\nLanguage: {project.get('language', 'English')}\n{len(project.get('chapters', []))} Chapters")
        meta_run.font.size = Pt(12)
        
        # Add author bio if present
        if project.get('author_bio'):
            doc.add_paragraph()
            author_paragraph = doc.add_paragraph()
            author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_paragraph.add_run(f"By: {project['author_bio']}")
            author_run.font.size = Pt(14)
            author_run.font.italic = True
        
        doc.add_page_break()
        
        # Table of contents
        toc_heading = doc.add_paragraph()
        toc_run = toc_heading.add_run("Table of Contents")
        toc_run.font.size = Pt(18)
        toc_run.font.bold = True
        doc.add_paragraph()
        
        # Add chapters to TOC
        for i, chapter in enumerate(project.get('chapters', []), 1):
            if chapter.get('status') == 'completed' and chapter.get('content'):
                toc_entry = doc.add_paragraph(f"Chapter {i}: {chapter['title']}")
                toc_entry.style = 'List Number'
        
        doc.add_page_break()
        
        # Add chapters
        for i, chapter in enumerate(project.get('chapters', []), 1):
            if chapter.get('status') == 'completed' and chapter.get('content'):
                # Chapter title
                chapter_heading = doc.add_paragraph()
                chapter_run = chapter_heading.add_run(f"Chapter {i}: {chapter['title']}")
                chapter_run.font.size = Pt(16)
                chapter_run.font.bold = True
                doc.add_paragraph()
                
                # Chapter content
                content = chapter['content']
                
                # Clean markdown formatting
                content = content.replace('**', '')
                content = content.replace('*', '')
                content = content.replace('#', '')
                
                # Split content into paragraphs
                paragraphs = content.split('\n\n')
                for paragraph_text in paragraphs:
                    if paragraph_text.strip():
                        paragraph = doc.add_paragraph(paragraph_text.strip())
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                doc.add_page_break()
        
        # Save DOCX file
        docx_filename = f"{project['name']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        docx_path = os.path.join(EXPORTS_FOLDER, docx_filename)
        doc.save(docx_path)
        
        return send_file(docx_path, as_attachment=True, download_name=docx_filename)
        
    except Exception as e:
        flash(f'Error exporting DOCX: {str(e)}', 'error')
        return redirect(url_for('project_view', project_id=project_id))

@app.route('/save_chapter/<chapter_id>', methods=['POST'])
def save_chapter_content(chapter_id):
    """Save edited chapter content and title via API"""
    try:
        data = request.get_json()
        content = data.get('content', '')
        title = data.get('title', '')
        
        # Find the project containing this chapter
        for filename in os.listdir(PROJECTS_FOLDER):
            if filename.endswith('.json'):
                project_file = os.path.join(PROJECTS_FOLDER, filename)
                with open(project_file, 'r') as f:
                    project = json.load(f)
                
                # Find and update the chapter
                for chapter in project.get('chapters', []):
                    if chapter['id'] == chapter_id:
                        chapter['content'] = content
                        if title:
                            chapter['title'] = title
                        project['last_modified'] = datetime.now().isoformat()
                        
                        with open(project_file, 'w') as f:
                            json.dump(project, f, indent=2)
                        
                        return jsonify({'success': True})
        
        return jsonify({'success': False, 'message': 'Chapter not found'}), 404
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/enhance_author_bio', methods=['POST'])
def enhance_author_bio():
    """Enhance author bio using AI"""
    try:
        config = load_config()
        data = request.get_json()
        
        author_info = data.get('author_info', '').strip()
        bio_style = data.get('bio_style', 'professional')
        bio_length = data.get('bio_length', 'medium')
        
        if not author_info:
            return jsonify({'success': False, 'message': 'Author information is required'})
        
        # Check if AI provider is configured
        ai_provider = config.get('ai_provider', 'openrouter')
        
        if ai_provider == 'gemini':
            if not config.get('gemini_api_key'):
                return jsonify({'success': False, 'message': 'Gemini API key not configured. Please add it in settings.'})
            
            enhanced_bio = generate_author_bio_gemini(author_info, bio_style, bio_length, config)
        else:
            if not config.get('openrouter_api_key'):
                return jsonify({'success': False, 'message': 'OpenRouter API key not configured. Please add it in settings.'})
            
            enhanced_bio = generate_author_bio_openrouter(author_info, bio_style, bio_length, config)
        
        return jsonify({'success': True, 'enhanced_bio': enhanced_bio})
        
    except Exception as e:
        logging.error(f"Error enhancing author bio: {str(e)}")
        return jsonify({'success': False, 'message': f'Error generating bio: {str(e)}'})

def generate_author_bio_gemini(author_info, bio_style, bio_length, config):
    """Generate author bio using Gemini AI"""
    try:
        import os
        from google import genai
        
        # Set up Gemini client
        client = genai.Client(api_key=config['gemini_api_key'])
        
        # Create style-specific prompt
        style_prompts = {
            'professional': 'Write in a professional, formal tone suitable for academic or business contexts',
            'conversational': 'Write in a warm, engaging, conversational tone that connects with readers',
            'academic': 'Write in a scholarly, academic tone with emphasis on credentials and expertise',
            'creative': 'Write in an expressive, creative tone that showcases personality and uniqueness'
        }
        
        length_instructions = {
            'short': 'Keep it to 2-3 sentences maximum',
            'medium': 'Write 1 paragraph (4-6 sentences)',
            'long': 'Write 2-3 paragraphs with comprehensive details'
        }
        
        prompt = f"""You are a professional biography writer. Create a compelling author biography based on the following information.

AUTHOR INFORMATION:
{author_info}

STYLE REQUIREMENTS:
- {style_prompts.get(bio_style, style_prompts['professional'])}
- {length_instructions.get(bio_length, length_instructions['medium'])}
- Focus on credibility, expertise, and what makes this author qualified to write
- Include relevant achievements, background, and experience
- Make it engaging and professional
- Write in third person
- Do not include any markdown formatting or special characters

Create the author biography now:"""

        response = client.models.generate_content(
            model=config.get('gemini_model', 'gemini-1.5-flash'),
            contents=prompt
        )
        
        if response.text:
            return response.text.strip()
        else:
            raise Exception("No response from Gemini API")
            
    except Exception as e:
        raise Exception(f"Gemini API error: {str(e)}")

def generate_author_bio_openrouter(author_info, bio_style, bio_length, config):
    """Generate author bio using OpenRouter API"""
    try:
        # Create style-specific prompt
        style_prompts = {
            'professional': 'Write in a professional, formal tone suitable for academic or business contexts',
            'conversational': 'Write in a warm, engaging, conversational tone that connects with readers',
            'academic': 'Write in a scholarly, academic tone with emphasis on credentials and expertise',
            'creative': 'Write in an expressive, creative tone that showcases personality and uniqueness'
        }
        
        length_instructions = {
            'short': 'Keep it to 2-3 sentences maximum',
            'medium': 'Write 1 paragraph (4-6 sentences)',
            'long': 'Write 2-3 paragraphs with comprehensive details'
        }
        
        prompt = f"""You are a professional biography writer. Create a compelling author biography based on the following information.

AUTHOR INFORMATION:
{author_info}

STYLE REQUIREMENTS:
- {style_prompts.get(bio_style, style_prompts['professional'])}
- {length_instructions.get(bio_length, length_instructions['medium'])}
- Focus on credibility, expertise, and what makes this author qualified to write
- Include relevant achievements, background, and experience
- Make it engaging and professional
- Write in third person
- Do not include any markdown formatting or special characters

Create the author biography now:"""

        headers = {
            "Authorization": f"Bearer {config['openrouter_api_key']}",
            "Content-Type": "application/json",
        }
        
        data = {
            "model": config.get('selected_model', 'meta-llama/llama-3.2-3b-instruct:free'),
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7,
            "max_tokens": 500
        }
        
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        
        if response.status_code == 200:
            result = response.json()
            if result.get('choices') and len(result['choices']) > 0:
                return result['choices'][0]['message']['content'].strip()
            else:
                raise Exception("No content generated")
        else:
            raise Exception(f"API error: {response.status_code} - {response.text}")
            
    except Exception as e:
        raise Exception(f"OpenRouter API error: {str(e)}")

@app.route('/check_generation_status/<project_id>')
def check_generation_status(project_id):
    try:
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        return jsonify({
            'status': project.get('generation_status', 'unknown'),
            'chapters': len(project.get('chapters', [])),
            'completed_chapters': len([c for c in project.get('chapters', []) if c.get('status') == 'completed'])
        })
    except:
        return jsonify({'status': 'error'})

@app.route('/api/check_ai_status')
def api_check_ai_status():
    """Check AI provider status and configuration"""
    try:
        config = load_config()
        
        # Check which AI provider is configured
        ai_provider = config.get('ai_provider', 'openrouter')
        is_ready = False
        status = "Not configured"
        detailed_status = "AI not configured"
        
        if ai_provider == 'gemini':
            api_key = config.get('gemini_api_key', '')
            if api_key:
                is_ready = True
                status = "Gemini Ready"
                detailed_status = f"Gemini AI ({config.get('gemini_model', 'gemini-1.5-flash')})"
            else:
                detailed_status = "Gemini API key missing"
        else:  # openrouter
            api_key = config.get('openrouter_api_key', '')
            if api_key:
                is_ready = True
                status = "OpenRouter Ready"
                detailed_status = f"OpenRouter ({config.get('selected_model', 'auto')})"
            else:
                detailed_status = "OpenRouter API key missing"
        
        return jsonify({
            'is_ready': is_ready,
            'status': status,
            'detailed_status': detailed_status,
            'provider': ai_provider
        })
    except Exception as e:
        logging.error(f"Error checking AI status: {e}")
        return jsonify({
            'is_ready': False,
            'status': "Error",
            'detailed_status': "Status check failed",
            'provider': 'unknown'
        }), 500

@app.route('/api/test_ai_connection')
def api_test_ai_connection():
    """Test AI connection with a simple request"""
    try:
        config = load_config()
        ai_provider = config.get('ai_provider', 'openrouter')
        
        if ai_provider == 'gemini':
            api_key = config.get('gemini_api_key', '')
            if not api_key:
                return jsonify({'success': False, 'message': 'Gemini API key not configured'})
            
            # Simple test request to Gemini using new library
            try:
                from google import genai
                from google.genai import types
            except ImportError:
                return jsonify({'success': False, 'message': 'Gemini library not installed. Please install google-genai.'})
            
            try:
                client = genai.Client(api_key=api_key)
                response = client.models.generate_content(
                    model=config.get('gemini_model', 'gemini-2.5-flash'),
                    contents="Hello, this is a connection test."
                )
                if response.text:
                    return jsonify({'success': True, 'message': 'Gemini connection successful'})
                else:
                    return jsonify({'success': False, 'message': 'Gemini test failed: No response'})
            except Exception as e:
                return jsonify({'success': False, 'message': f'Gemini test failed: {str(e)}'})
                
        else:  # openrouter
            api_key = config.get('openrouter_api_key', '')
            if not api_key:
                return jsonify({'success': False, 'message': 'OpenRouter API key not configured'})
            
            # Simple test request to OpenRouter
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }
            
            data = {
                "model": config.get('selected_model', 'meta-llama/llama-3.2-3b-instruct:free'),
                "messages": [{"role": "user", "content": "Hello, this is a connection test."}],
                "max_tokens": 50
            }
            
            try:
                response = requests.post(
                    "https://openrouter.ai/api/v1/chat/completions",
                    headers=headers,
                    json=data,
                    timeout=30
                )
                
                if response.status_code == 200:
                    return jsonify({'success': True, 'message': 'OpenRouter connection successful'})
                else:
                    return jsonify({'success': False, 'message': f'OpenRouter test failed: {response.status_code}'})
            except Exception as e:
                return jsonify({'success': False, 'message': f'OpenRouter test failed: {str(e)}'})
                
    except Exception as e:
        logging.error(f"Error testing AI connection: {e}")
        return jsonify({'success': False, 'message': f'Connection test error: {str(e)}'})



# Session storage for standalone generation
standalone_sessions = {}

@app.route('/standalone_generation')
def standalone_generation():
    """Standalone AI book generation page"""
    return render_template('standalone_generation.html')

@app.route('/api/standalone_generation', methods=['POST'])
def api_standalone_generation():
    """Start standalone book generation"""
    try:
        config = load_config()
        if not config.get('license_activated', False):
            return jsonify({'success': False, 'message': 'License not activated'})
            
        data = request.get_json()
        session_id = str(uuid.uuid4())
        
        # Store session data
        standalone_sessions[session_id] = {
            'status': 'starting',
            'currentStep': 'enhancing',
            'completedSteps': [],
            'book': {
                'title': '',
                'description': data.get('topic', ''),
                'authorBio': data.get('authorBio', ''),
                'language': data.get('language', 'English'),
                'chapters': []
            },
            'params': data,
            'error': None
        }
        
        # Start background generation
        thread = threading.Thread(target=standalone_generation_background, 
                                args=(session_id, data, config))
        thread.daemon = True
        thread.start()
        
        return jsonify({'success': True, 'session_id': session_id})
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/standalone_status/<session_id>')
def api_standalone_status(session_id):
    """Get standalone generation status"""
    if session_id not in standalone_sessions:
        return jsonify({'status': 'not_found'})
        
    return jsonify(standalone_sessions[session_id])

@app.route('/api/standalone_export_pdf', methods=['POST'])
def api_standalone_export_pdf():
    """Export standalone book as PDF"""
    try:
        book_data = json.loads(request.form.get('book_data', '{}'))
        
        # Generate HTML for PDF
        html_content = render_template('book_export.html', project=book_data)
        
        # Convert to PDF
        pdf_filename = f"{book_data['title']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        pdf_path = os.path.join(EXPORTS_FOLDER, pdf_filename)
        
        weasyprint.HTML(string=html_content, base_url=request.url_root).write_pdf(pdf_path)
        
        return send_file(pdf_path, as_attachment=True, download_name=pdf_filename)
        
    except Exception as e:
        flash(f'Error exporting PDF: {str(e)}', 'error')
        return redirect(url_for('standalone_generation'))

@app.route('/api/standalone_export_docx', methods=['POST'])
def api_standalone_export_docx():
    """Export standalone book as DOCX"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        book_data = json.loads(request.form.get('book_data', '{}'))
        
        # Create DOCX document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Cover page with image if available
        if book_data.get('cover_image'):
            cover_image_path = os.path.join(UPLOAD_FOLDER, book_data['cover_image'])
            if os.path.exists(cover_image_path):
                # Add cover image
                cover_paragraph = doc.add_paragraph()
                cover_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cover_run = cover_paragraph.add_run()
                try:
                    # Add image with full width
                    cover_run.add_picture(cover_image_path, width=Inches(6))
                except Exception as e:
                    # If image fails, fall back to text cover
                    pass
        
        # Title overlay or regular title
        doc.add_paragraph()  # Space
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(book_data['title'])
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        
        # Subtitle/description
        if book_data.get('description'):
            topic_paragraph = doc.add_paragraph()
            topic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            topic_run = topic_paragraph.add_run(book_data['description'])
            topic_run.font.size = Pt(16)
            topic_run.font.italic = True
        
        doc.add_paragraph()  # Empty line
        
        # Book metadata
        meta_paragraph = doc.add_paragraph()
        meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_paragraph.add_run(f"Generated with BookGenPro\nLanguage: {book_data.get('language', 'English')}\n{len(book_data.get('chapters', []))} Chapters")
        meta_run.font.size = Pt(12)
        
        # Author bio
        if book_data.get('authorBio'):
            doc.add_paragraph()
            author_paragraph = doc.add_paragraph()
            author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_paragraph.add_run(f"By: {book_data['authorBio']}")
            author_run.font.size = Pt(14)
            author_run.font.italic = True
        
        doc.add_page_break()
        
        # Table of contents
        toc_heading = doc.add_paragraph()
        toc_run = toc_heading.add_run("Table of Contents")
        toc_run.font.size = Pt(18)
        toc_run.font.bold = True
        doc.add_paragraph()
        
        for i, chapter in enumerate(book_data.get('chapters', []), 1):
            toc_entry = doc.add_paragraph(f"Chapter {i}: {chapter['title']}")
            toc_entry.style = 'List Number'
        
        doc.add_page_break()
        
        # Chapters
        for i, chapter in enumerate(book_data.get('chapters', []), 1):
            chapter_heading = doc.add_paragraph()
            chapter_run = chapter_heading.add_run(f"Chapter {i}: {chapter['title']}")
            chapter_run.font.size = Pt(16)
            chapter_run.font.bold = True
            doc.add_paragraph()
            
            content = chapter['content']
            content = content.replace('**', '').replace('*', '').replace('#', '')
            
            paragraphs = content.split('\n\n')
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    paragraph = doc.add_paragraph(paragraph_text.strip())
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            doc.add_page_break()
        
        # Save DOCX
        docx_filename = f"{book_data['title']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        docx_path = os.path.join(EXPORTS_FOLDER, docx_filename)
        doc.save(docx_path)
        
        return send_file(docx_path, as_attachment=True, download_name=docx_filename)
        
    except Exception as e:
        flash(f'Error exporting DOCX: {str(e)}', 'error')
        return redirect(url_for('standalone_generation'))

def standalone_generation_background(session_id, params, config):
    """Background process for standalone generation"""
    try:
        session = standalone_sessions[session_id]
        
        # Step 1: Enhance description
        session['status'] = 'generating'
        session['currentStep'] = 'enhancing'
        
        ai_provider = config.get('ai_provider', 'openrouter')
        enhanced_description = enhance_description_simple(params['topic'], params['language'], config)
        
        session['book']['description'] = enhanced_description
        session['completedSteps'].append('enhancing')
        
        # Step 2: Generate chapter titles
        session['currentStep'] = 'titles'
        
        titles = generate_chapter_titles_simple(enhanced_description, params['language'], int(params['chapterCount']), config)
        
        # Generate book title from first title or description
        if titles:
            book_title = generate_book_title_from_description(enhanced_description, config)
            session['book']['title'] = book_title
            
            # Initialize chapters
            for i, title in enumerate(titles):
                session['book']['chapters'].append({
                    'id': str(i+1),
                    'title': title,
                    'content': '',
                    'status': 'pending'
                })
        
        session['completedSteps'].append('titles')
        
        # Step 3: Generate content if full generation requested
        if params['type'] == 'full':
            session['currentStep'] = 'content'
            
            total_chapters = len(session['book']['chapters'])
            for i, chapter in enumerate(session['book']['chapters']):
                session['progress'] = f"({i+1}/{total_chapters})"
                
                content = generate_chapter_content_simple(
                    enhanced_description, chapter['title'], params['language'], config
                )
                
                chapter['content'] = content
                chapter['status'] = 'completed'
        
        session['completedSteps'].append('content')
        session['status'] = 'completed'
        
        # Save to main project system
        save_standalone_to_project(session, params)
        
    except Exception as e:
        if 'session' in locals():
            session['status'] = 'error'
            session['error'] = str(e)

def save_standalone_to_project(session, params):
    """Save standalone generation to main project system"""
    try:
        project_id = str(uuid.uuid4())
        
        project_data = {
            'id': project_id,
            'name': session['book']['title'],
            'topic': session['book']['description'],
            'language': session['book']['language'],
            'chapters': session['book']['chapters'],
            'status': 'completed',
            'author_bio': session['book']['authorBio'],
            'created_at': datetime.now().isoformat(),
            'last_modified': datetime.now().isoformat(),
            'generation_type': 'ai_standalone'
        }
        
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'w') as f:
            json.dump(project_data, f, indent=2)
            
        if 'session' in locals():
            session['project_id'] = project_id
        
    except Exception as e:
        print(f"Error saving standalone project: {e}")

def enhance_description_simple(description, language, config):
    """Simple description enhancement"""
    ai_provider = config.get('ai_provider', 'openrouter')
    
    prompt = f"""Enhance this book description for better AI generation in {language}:

{description}

Make it more detailed and comprehensive while maintaining the core concept. Add context that would help generate better chapter titles and content."""

    try:
        if ai_provider == 'gemini':
            success, result = generate_with_gemini(prompt, config.get('gemini_api_key', ''), config.get('gemini_model', 'gemini-2.5-flash'))
            if success:
                return result.strip()
            else:
                logging.error(f"Gemini failed in enhance_description_simple: {result}")
        
        # OpenRouter fallback
        headers = {
            "Authorization": f"Bearer {config['openrouter_api_key']}",
            "Content-Type": "application/json",
        }
        
        data = {
            "model": config.get('selected_model', 'meta-llama/llama-3.2-3b-instruct:free'),
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7,
            "max_tokens": 500
        }
        
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        
        if response.status_code == 200:
            result = response.json()
            if result.get('choices') and len(result['choices']) > 0:
                return result['choices'][0]['message']['content'].strip()
        
    except Exception as e:
        print(f"Enhancement error: {e}")
    
    return description

def generate_chapter_titles_simple(description, language, chapter_count, config):
    """Generate chapter titles"""
    ai_provider = config.get('ai_provider', 'openrouter')
    
    prompt = f"""Create {chapter_count} engaging chapter titles in {language} for a book about:

{description}

Requirements:
- Each title should be compelling and descriptive
- Arrange them in logical order
- Make them flow well together
- Return only the titles, one per line
- No numbering or formatting

Chapter titles:"""

    try:
        if ai_provider == 'gemini':
            success, result = generate_with_gemini(prompt, config.get('gemini_api_key', ''), config.get('gemini_model', 'gemini-2.5-flash'))
            if success:
                titles = result.strip().split('\n')
                return [title.strip() for title in titles if title.strip()]
            else:
                logging.error(f"Gemini failed in generate_chapter_titles_simple: {result}")
        
        # OpenRouter fallback
        headers = {
            "Authorization": f"Bearer {config['openrouter_api_key']}",
            "Content-Type": "application/json",
        }
        
        data = {
            "model": config.get('selected_model', 'meta-llama/llama-3.2-3b-instruct:free'),
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7,
            "max_tokens": 800
        }
        
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        
        if response.status_code == 200:
            result = response.json()
            if result.get('choices') and len(result['choices']) > 0:
                titles = result['choices'][0]['message']['content'].strip().split('\n')
                return [title.strip() for title in titles if title.strip()]
        
    except Exception as e:
        print(f"Title generation error: {e}")
    
    return [f"Chapter {i+1}" for i in range(chapter_count)]

def generate_chapter_content_simple(description, chapter_title, language, config):
    """Generate chapter content"""
    ai_provider = config.get('ai_provider', 'openrouter')
    
    prompt = f"""Write comprehensive chapter content in {language} for:

Book Description: {description}
Chapter Title: {chapter_title}

Requirements:
- Write 800-1200 words
- Make it engaging and informative
- Use clear paragraphs
- Match the book's overall theme
- Don't include chapter title in the content
- Write in a natural, flowing style

Chapter content:"""

    try:
        if ai_provider == 'gemini':
            success, result = generate_with_gemini(prompt, config.get('gemini_api_key', ''), config.get('gemini_model', 'gemini-2.5-flash'))
            if success:
                return result.strip()
            else:
                logging.error(f"Gemini failed in generate_chapter_content_simple: {result}")
        
        # OpenRouter fallback
        headers = {
            "Authorization": f"Bearer {config['openrouter_api_key']}",
            "Content-Type": "application/json",
        }
        
        data = {
            "model": config.get('selected_model', 'meta-llama/llama-3.2-3b-instruct:free'),
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7,
            "max_tokens": 2000
        }
        
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        
        if response.status_code == 200:
            result = response.json()
            if result.get('choices') and len(result['choices']) > 0:
                return result['choices'][0]['message']['content'].strip()
        
    except Exception as e:
        print(f"Content generation error: {e}")
    
    return f"Content for {chapter_title} will be generated here."

def generate_book_title_from_description(description, config):
    """Generate book title from description"""
    try:
        ai_provider = config.get('ai_provider', 'openrouter')
        
        prompt = f"""Create a compelling book title based on this description:
        
{description}

Requirements:
- Make it catchy and professional
- Keep it under 60 characters
- Make it marketable and appealing
- Don't use quotes around the title
- Just return the title only

Title:"""

        if ai_provider == 'gemini':
            success, result = generate_with_gemini(prompt, config.get('gemini_api_key', ''), config.get('gemini_model', 'gemini-2.5-flash'))
            if success:
                return result.strip()
            else:
                logging.error(f"Gemini failed in generate_book_title_from_description: {result}")
                return "Generated Book"
        else:
            headers = {
                "Authorization": f"Bearer {config['openrouter_api_key']}",
                "Content-Type": "application/json",
            }
            
            data = {
                "model": config.get('selected_model', 'meta-llama/llama-3.2-3b-instruct:free'),
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.7,
                "max_tokens": 100
            }
            
            response = requests.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers=headers,
                json=data,
                timeout=60
            )
            
            if response.status_code == 200:
                result = response.json()
                if result.get('choices') and len(result['choices']) > 0:
                    return result['choices'][0]['message']['content'].strip()
            
            return "Generated Book"
            
    except Exception as e:
        return "Generated Book"

# ===== AI ENHANCEMENT API ROUTES =====

@app.route('/api/enhance_title/<project_id>', methods=['POST'])
def enhance_title(project_id):
    """Enhance project title using AI"""
    try:
        config = load_config()
        if not config.get('license_activated', False):
            return jsonify({'success': False, 'message': 'License not activated'})
        
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        current_title = project.get('name', '')
        topic = project.get('topic', '')
        language = project.get('language', 'English')
        
        ai_provider = config.get('ai_provider', 'openrouter')
        
        # Create enhancement prompt
        prompt = f"""Enhance this book title to make it more compelling and professional:

Current Title: "{current_title}"
Book Topic: "{topic}"
Language: {language}

Requirements:
1. Keep the core meaning intact
2. Make it more engaging and marketable
3. Ensure it accurately reflects the book's content
4. Make it memorable and professional
5. Optimal length for a book title (not too long)

Return only the enhanced title, nothing else."""

        # Generate enhanced title
        if ai_provider == 'gemini':
            success, enhanced_title = generate_with_gemini(prompt, config.get('gemini_api_key', ''), config.get('gemini_model', 'gemini-2.5-flash'))
            if not success:
                return jsonify({'success': False, 'message': 'AI generation failed'})
        else:
            success, enhanced_title = generate_content(prompt, config)
            if not success:
                return jsonify({'success': False, 'message': 'AI generation failed'})
        
        # Update project
        project['name'] = enhanced_title.strip()
        project['last_modified'] = datetime.now().isoformat()
        
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
        
        return jsonify({'success': True, 'enhanced_title': enhanced_title.strip()})
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/enhance_description/<project_id>', methods=['POST'])
def enhance_description(project_id):
    """Enhance project description using AI"""
    try:
        config = load_config()
        if not config.get('license_activated', False):
            return jsonify({'success': False, 'message': 'License not activated'})
        
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        current_description = project.get('topic', '')
        title = project.get('name', '')
        language = project.get('language', 'English')
        
        ai_provider = config.get('ai_provider', 'openrouter')
        
        # Create enhancement prompt
        prompt = f"""Enhance this book description to be more detailed and compelling:

Book Title: "{title}"
Current Description: "{current_description}"
Language: {language}

Requirements:
1. Expand on the core concept while maintaining focus
2. Add context and depth to make it more comprehensive
3. Make it engaging for potential readers
4. Include key topics that should be covered
5. Maintain professional tone
6. Keep it clear and well-structured

Return the enhanced description, nothing else."""

        # Generate enhanced description
        if ai_provider == 'gemini':
            success, enhanced_description = generate_with_gemini(prompt, config.get('gemini_api_key', ''), config.get('gemini_model', 'gemini-2.5-flash'))
            if not success:
                return jsonify({'success': False, 'message': 'AI generation failed'})
        else:
            success, enhanced_description = generate_content(prompt, config)
            if not success:
                return jsonify({'success': False, 'message': 'AI generation failed'})
        
        # Update project
        project['topic'] = enhanced_description.strip()
        project['last_modified'] = datetime.now().isoformat()
        
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
        
        return jsonify({'success': True, 'enhanced_description': enhanced_description.strip()})
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/update_chapter_title', methods=['POST'])
def update_chapter_title():
    """Update a specific chapter title"""
    try:
        config = load_config()
        if not config.get('license_activated', False):
            return jsonify({'success': False, 'message': 'License not activated'})
        
        data = request.get_json()
        project_id = data.get('project_id')
        chapter_id = data.get('chapter_id')
        new_title = data.get('title')
        
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        # Find and update the chapter
        for chapter in project.get('chapters', []):
            if chapter.get('id') == chapter_id:
                chapter['title'] = new_title
                break
        
        project['last_modified'] = datetime.now().isoformat()
        
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
        
        return jsonify({'success': True, 'message': 'Chapter title updated successfully'})
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/enhance_chapter_title/<project_id>/<chapter_id>', methods=['POST'])
def enhance_chapter_title(project_id, chapter_id):
    """Enhance a specific chapter title using AI"""
    try:
        config = load_config()
        if not config.get('license_activated', False):
            return jsonify({'success': False, 'message': 'License not activated'})
        
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        # Find the chapter
        target_chapter = None
        for chapter in project.get('chapters', []):
            if chapter.get('id') == chapter_id:
                target_chapter = chapter
                break
        
        if not target_chapter:
            return jsonify({'success': False, 'message': 'Chapter not found'})
        
        current_title = target_chapter.get('title', '')
        book_title = project.get('name', '')
        book_topic = project.get('topic', '')
        chapter_content = target_chapter.get('content', '')[:500]  # First 500 chars
        language = project.get('language', 'English')
        
        ai_provider = config.get('ai_provider', 'openrouter')
        
        # Create enhancement prompt
        prompt = f"""Enhance this chapter title to be more compelling and descriptive:

Book Title: "{book_title}"
Book Topic: "{book_topic}"
Current Chapter Title: "{current_title}"
Chapter Content Preview: "{chapter_content}..."
Language: {language}

Requirements:
1. Make the title more engaging and specific
2. Ensure it accurately reflects the chapter content
3. Keep it professional and appropriate for the book
4. Make it clear what readers will learn from this chapter
5. Maintain appropriate length for a chapter title
6. Ensure it fits well with the overall book structure

Return only the enhanced chapter title, nothing else."""

        # Generate enhanced title
        if ai_provider == 'gemini':
            success, enhanced_title = generate_with_gemini(prompt, config.get('gemini_api_key', ''), config.get('gemini_model', 'gemini-2.5-flash'))
            if not success:
                return jsonify({'success': False, 'message': 'AI generation failed'})
        else:
            success, enhanced_title = generate_content(prompt, config)
            if not success:
                return jsonify({'success': False, 'message': 'AI generation failed'})
        
        # Update chapter title
        target_chapter['title'] = enhanced_title.strip()
        project['last_modified'] = datetime.now().isoformat()
        
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
        
        return jsonify({'success': True, 'enhanced_title': enhanced_title.strip()})
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/enhance_chapter_content/<project_id>/<chapter_id>', methods=['POST'])
def enhance_chapter_content(project_id, chapter_id):
    """Enhance a specific chapter content using AI"""
    try:
        config = load_config()
        if not config.get('license_activated', False):
            return jsonify({'success': False, 'message': 'License not activated'})
        
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        # Find the chapter
        target_chapter = None
        for chapter in project.get('chapters', []):
            if chapter.get('id') == chapter_id:
                target_chapter = chapter
                break
        
        if not target_chapter:
            return jsonify({'success': False, 'message': 'Chapter not found'})
        
        current_content = target_chapter.get('content', '')
        chapter_title = target_chapter.get('title', '')
        book_title = project.get('name', '')
        book_topic = project.get('topic', '')
        language = project.get('language', 'English')
        
        ai_provider = config.get('ai_provider', 'openrouter')
        
        # Create enhancement prompt
        prompt = f"""Enhance and improve this chapter content while maintaining its core structure and meaning:

Book Title: "{book_title}"
Book Topic: "{book_topic}"
Chapter Title: "{chapter_title}"
Language: {language}

Current Content:
{current_content}

Requirements:
1. Improve clarity and flow of the writing
2. Enhance explanations and add helpful details where appropriate
3. Maintain the original structure and key points
4. Improve transitions between paragraphs
5. Ensure professional, engaging tone
6. Add depth without changing the fundamental content
7. Keep the same approximate length
8. Maintain consistency with the book's overall theme

Return the enhanced content only, preserving paragraph structure."""

        # Generate enhanced content
        if ai_provider == 'gemini':
            success, enhanced_content = generate_with_gemini(prompt, config.get('gemini_api_key', ''), config.get('gemini_model', 'gemini-2.5-flash'))
            if not success:
                return jsonify({'success': False, 'message': 'AI generation failed'})
        else:
            success, enhanced_content = generate_content(prompt, config)
            if not success:
                return jsonify({'success': False, 'message': 'AI generation failed'})
        
        # Update chapter content
        target_chapter['content'] = enhanced_content.strip()
        target_chapter['status'] = 'completed'
        project['last_modified'] = datetime.now().isoformat()
        
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
        
        return jsonify({'success': True, 'message': 'Chapter content enhanced successfully'})
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/enhance_all_chapter_titles/<project_id>', methods=['POST'])
def enhance_all_chapter_titles(project_id):
    """Enhance all chapter titles using AI"""
    try:
        config = load_config()
        if not config.get('license_activated', False):
            return jsonify({'success': False, 'message': 'License not activated'})
        
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        chapters = project.get('chapters', [])
        if not chapters:
            return jsonify({'success': False, 'message': 'No chapters found'})
        
        book_title = project.get('name', '')
        book_topic = project.get('topic', '')
        language = project.get('language', 'English')
        
        ai_provider = config.get('ai_provider', 'openrouter')
        
        # Collect all current titles
        current_titles = [f"Chapter {ch.get('number', i+1)}: {ch.get('title', '')}" for i, ch in enumerate(chapters)]
        
        # Create enhancement prompt for all titles
        prompt = f"""Enhance ALL these chapter titles to be more compelling and descriptive:

Book Title: "{book_title}"
Book Topic: "{book_topic}"
Language: {language}

Current Chapter Titles:
{chr(10).join(current_titles)}

Requirements:
1. Enhance each title to be more engaging and specific
2. Ensure titles work well together as a cohesive sequence
3. Make each title clearly indicate what the chapter covers
4. Keep professional tone appropriate for the book
5. Maintain logical progression between chapters
6. Ensure titles are descriptive but not too long

Return the enhanced titles in the same order, one per line, in format "Chapter X: [Enhanced Title]"."""

        # Generate enhanced titles
        if ai_provider == 'gemini':
            success, enhanced_titles_text = generate_with_gemini(prompt, config.get('gemini_api_key', ''), config.get('gemini_model', 'gemini-2.5-flash'))
            if not success:
                return jsonify({'success': False, 'message': 'AI generation failed'})
        else:
            success, enhanced_titles_text = generate_content(prompt, config)
            if not success:
                return jsonify({'success': False, 'message': 'AI generation failed'})
        
        # Parse enhanced titles
        enhanced_lines = enhanced_titles_text.strip().split('\n')
        enhanced_titles = []
        for line in enhanced_lines:
            if ':' in line:
                # Extract title after "Chapter X: "
                title_part = line.split(':', 1)[1].strip()
                enhanced_titles.append(title_part)
        
        # Update all chapter titles
        for i, chapter in enumerate(chapters):
            if i < len(enhanced_titles):
                chapter['title'] = enhanced_titles[i]
        
        project['last_modified'] = datetime.now().isoformat()
        
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
        
        return jsonify({'success': True, 'message': f'Enhanced {len(enhanced_titles)} chapter titles successfully'})
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/regenerate_all_chapter_content/<project_id>', methods=['POST'])
def regenerate_all_chapter_content(project_id):
    """Regenerate all chapter content using AI"""
    try:
        config = load_config()
        if not config.get('license_activated', False):
            return jsonify({'success': False, 'message': 'License not activated'})
        
        project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
        with open(project_file, 'r') as f:
            project = json.load(f)
        
        chapters = project.get('chapters', [])
        if not chapters:
            return jsonify({'success': False, 'message': 'No chapters found'})
        
        # Start background regeneration
        thread = threading.Thread(target=regenerate_all_content_background, 
                                args=(project_id, project, config))
        thread.daemon = True
        thread.start()
        
        return jsonify({'success': True, 'message': 'Started regenerating all chapter content in background'})
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

def regenerate_all_content_background(project_id, project, config):
    """Background task to regenerate all chapter content"""
    project_file = os.path.join(PROJECTS_FOLDER, f"{project_id}.json")
    
    try:
        book_title = project.get('name', '')
        book_topic = project.get('topic', '')
        language = project.get('language', 'English')
        ai_provider = config.get('ai_provider', 'openrouter')
        
        # Update status
        project['generation_status'] = 'regenerating_all_content'
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
        
        # Regenerate each chapter
        for i, chapter in enumerate(project.get('chapters', [])):
            chapter['status'] = 'generating'
            with open(project_file, 'w') as f:
                json.dump(project, f, indent=2)
            
            chapter_title = chapter.get('title', '')
            
            # Create content generation prompt
            prompt = f"""Generate comprehensive content for this chapter:

Book Title: "{book_title}"
Book Topic: "{book_topic}"
Chapter Title: "{chapter_title}"
Language: {language}

Requirements:
1. Write detailed, informative content appropriate for this chapter
2. Ensure content aligns with the book's overall theme
3. Use professional, engaging writing style
4. Include practical examples or insights where relevant
5. Structure content with clear paragraphs
6. Aim for substantial content (500-1000 words)
7. Make it valuable and educational for readers

Generate the complete chapter content:"""

            # Generate content
            if ai_provider == 'gemini':
                success, content = generate_with_gemini(prompt, config.get('gemini_api_key', ''), config.get('gemini_model', 'gemini-2.5-flash'))
            else:
                success, content = generate_content(prompt, config)
            
            if success:
                chapter['content'] = content.strip()
                chapter['status'] = 'completed'
            else:
                chapter['status'] = 'error'
            
            with open(project_file, 'w') as f:
                json.dump(project, f, indent=2)
        
        # Mark as completed
        project['generation_status'] = 'completed'
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)
    
    except Exception as e:
        logging.error(f"Error in regenerate_all_content_background: {e}")
        project['generation_status'] = f'error: {str(e)}'
        with open(project_file, 'w') as f:
            json.dump(project, f, indent=2)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
